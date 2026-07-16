import streamlit as st
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import urllib.parse

# Configuration de la page web
st.set_page_config(page_title="Annuaire - Club des Entreprises", page_icon="🏢", layout="centered")

st.title("📝 Annuaire des Entreprises")
st.subheader("Club des Entreprises de la Vallée de l'Isle")
st.markdown("Renseignez vos informations, déposez vos visuels, puis validez pour transmettre votre fiche.")

# --- FONCTION DE GÉNÉRATION WORD ---
def generate_word_doc(data, logo_files, portrait_file, illustr_files):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    title = doc.add_paragraph()
    title_run = title.add_run("FICHE ANNUAIRE ENTREPRISE")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(20)

    def add_section_header(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title_text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        pBrd = OxmlElement('w:pBrd')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'D3D3D3')
        pBrd.append(bottom)
        p._p.get_or_add_pPr().append(pBrd)

    def add_field(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        lbl_run = p.add_run(f"• {label} : ")
        lbl_run.font.bold = True
        val_run = p.add_run(value if value else "Non renseigné")
        if not value:
            val_run.font.italic = True
            val_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # SECTION VISUELS
    add_section_header("ÉLÉMENTS VISUELS")
    
    # Gestion de plusieurs logos
    if logo_files:
        doc.add_paragraph().add_run(f"• Logo(s) de l'entreprise ({len(logo_files)}) :").font.bold = True
        for logo in logo_files:
            doc.add_picture(logo, width=Inches(2.2))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
    else:
        add_field("Logo de l'entreprise", "Non fourni")

    # Gestion de la photo portrait
    if portrait_file is not None:
        p_port = doc.add_paragraph()
        p_port.paragraph_format.space_before = Pt(8)
        p_port.add_run("• Photo du dirigeant :").font.bold = True
        doc.add_picture(portrait_file, width=Inches(1.8))
    else:
        add_field("Photo du dirigeant", "Non fournie")

    # Gestion de plusieurs illustrations
    if illustr_files:
        p_ill = doc.add_paragraph()
        p_ill.paragraph_format.space_before = Pt(8)
        p_ill.add_run(f"• Visuels d'illustration ({len(illustr_files)}) :").font.bold = True
        for img in illustr_files:
            doc.add_picture(img, width=Inches(2.8))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
    else:
        add_field("Visuels d'illustration", "Aucun visuel fourni")

    # SECTIONS TEXTES
    add_section_header("1. CATÉGORIE DE L'ENTREPRISE")
    add_field("Secteur d'activité principal", data["categorie"])

    add_section_header("2. IDENTITÉ DE L'ENTREPRISE")
    add_field("Nom de l'entreprise (Raison sociale)", data["raison_sociale"])
    add_field("Nom et Prénom du Dirigeant", data["dirigeant"])
    add_field("Fonction", data["fonction"])
    add_field("Site Internet", data["site_web"])

    add_section_header("3. CHIFFRES CLÉS & STRUCTURE")
    add_field("Effectif", data["effectif"])
    add_field("Structure / Groupe", data["structure"])
    add_field("Données marquantes", data["donnees_marquantes"])

    add_section_header("4. DESCRIPTIF DE L'ACTIVITÉ & SPÉCIALITÉS")
    p = doc.add_paragraph()
    p.add_run(data["descriptif"] if data["descriptif"] else "Aucun descriptif fourni.")

    add_section_header("5. COORDONNÉES DE CONTACT")
    add_field("Adresse postale", data["adresse"])
    add_field("Téléphone", data["telephone"])
    add_field("Adresse e-mail", data["email"])
    add_field("Facebook", data["facebook"])
    add_field("Instagram", data["instagram"])
    add_field("LinkedIn", data["linkedin"])

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- INTERFACE DE SAISIE ---
st.markdown("### 📸 1. Vos fichiers visuels")

# Changement ici : "accept_multiple_files=True" est désormais actif pour le logo et les illustrations
logo_files = st.file_uploader("Logo(s) de l'entreprise (Déposez un ou plusieurs fichiers PNG, JPG)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
portrait_file = st.file_uploader("Photo portrait du dirigeant (Un seul fichier)", type=["png", "jpg", "jpeg"])
illustr_files = st.file_uploader("Visuels d'illustration (Locaux, équipe, produits... Déposez autant de fichiers que souhaité)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

st.markdown("---")
st.markdown("### 📝 2. Vos informations")

data = {}
data["categorie"] = st.text_input("Secteur d'activité principal (ex: Peinture, Imprimerie...)")

col1, col2 = st.columns(2)
with col1:
    data["raison_sociale"] = st.text_input("Nom de l'entreprise (Raison sociale)")
    data["dirigeant"] = st.text_input("Nom et Prénom du Dirigeant")
with col2:
    data["fonction"] = st.text_input("Fonction (ex: Gérant, Président...)")
    data["site_web"] = st.text_input("Site Internet")

data["effectif"] = st.text_input("Effectif de l'entreprise")
data["structure"] = st.text_input("Structure / Groupe (ex: Filiale de...)")
data["donnees_marquantes"] = st.text_input("Données marquantes optionnelles (ex: CA, Année de création)")
data["descriptif"] = st.text_area("Descriptif de l'activité & spécialités")
data["adresse"] = st.text_input("Adresse postale complète de l'établissement")

col3, col4 = st.columns(2)
with col3:
    data["telephone"] = st.text_input("Téléphone")
with col4:
    data["email"] = st.text_input("Adresse e-mail de contact")

st.markdown("**Réseaux sociaux**")
col5, col6, col7 = st.columns(3)
with col5: data["facebook"] = st.text_input("Facebook")
with col6: data["instagram"] = st.text_input("Instagram")
with col7: data["linkedin"] = st.text_input("LinkedIn")

st.markdown("---")

# --- ONGLET APERÇU ET ENVOI ---
if data["raison_sociale"]:
    st.markdown("### 👀 3. Aperçu avant validation")
    
    with st.container():
        st.info("Voici les informations telles qu'elles apparaîtront dans l'annuaire :")
        
        # Petit récapitulatif textuel des visuels ajoutés dans l'aperçu web
        nb_logos = len(logo_files) if logo_files else 0
        nb_ills = len(illustr_files) if illustr_files else 0
        has_port = "Oui" if portrait_file else "Non"
        
        st.markdown(f"""
        **📁 {data['raison_sociale'].upper()}**  
        *Secteur :* {data['categorie'] if data['categorie'] else 'Non renseigné'}  
        
        **🖼️ Statut des fichiers :** {nb_logos} logo(s) chargé(s) — Photo dirigeant : {has_port} — {nb_ills} image(s) d'illustration  
        
        **👤 Dirigeant :** {data['dirigeant']} ({data['fonction']})  
        *Site Web :* {data['site_web'] if data['site_web'] else 'Non renseigné'}  
        
        **📊 Chiffres & Structure :**  
        - Effectif : {data['effectif'] if data['effectif'] else 'Non renseigné'}  
        - Structure : {data['structure'] if data['structure'] else 'Non renseigné'}  
        - Repères : {data['donnees_marquantes'] if data['donnees_marquantes'] else 'Aucun'}  
        
        **📖 Descriptif :**  
        {data['descriptif'] if data['descriptif'] else 'Aucun descriptif rédigé.'}  
        
        **📞 Contacts :**  
        - Adresse : {data['adresse'] if data['adresse'] else 'Non renseignée'}  
        - Tél : {data['telephone'] if data['telephone'] else 'Non renseigné'}  
        - Email : {data['email'] if data['email'] else 'Non renseigné'}  
        """)

    st.markdown("### 🚀 4. Validation & Transmission")
    
    docx_file = generate_word_doc(data, logo_files, portrait_file, illustr_files)
    st.download_button(
        label="📥 1. Télécharger mon fichier Word (.docx)",
        data=docx_file,
        file_name=f"Fiche_{data['raison_sociale'].replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    sujet = f"Annuaire Club Entreprises - {data['raison_sociale']}"
    corps = f"Bonjour Jérôme,\n\nVoici les informations saisies pour l'annuaire du Club :\n\n" \
            f"Raison sociale : {data['raison_sociale']}\n" \
            f"Dirigeant : {data['dirigeant']} ({data['fonction']})\n" \
            f"Secteur : {data['categorie']}\n" \
            f"Fichiers : {nb_logos} logo(s) et {nb_ills} illustration(s) intégrés au Word.\n\n" \
            f"Pensez à glisser le fichier Word téléchargé ainsi que vos images originales en pièces jointes à cet e-mail."
    
    mail_link = f"mailto:j.gabuteau@invelac.fr?subject={urllib.parse.quote(sujet)}&body={urllib.parse.quote(corps)}"
    
    st.markdown(f"""
    <a href="{mail_link}" target="_blank">
        <button style="background-color: #2E7D32; color: white; font-weight: bold; padding: 12px 24px; border: none; border-radius: 4px; cursor: pointer; width: 100%;">
            ✉️ 2. Valider et envoyer à j.gabuteau@invelac.fr
        </button>
    </a>
    """, unsafe_allow_html=True)
    st.caption("Note : Le bouton d'envoi ouvrira votre logiciel de messagerie. Assurez-vous d'y joindre vos fichiers visuels avant de cliquer sur envoyer.")
else:
    st.warning("Veuillez saisir au moins le 'Nom de l'entreprise' pour débloquer l'aperçu et l'envoi.")